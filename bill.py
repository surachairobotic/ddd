from os import listdir
from os.path import isfile, join

import math
import os
import sys

from docx import Document
from docx.shared import Mm, Pt, Inches
from docx.enum.style import WD_STYLE_TYPE

#from docx2pdf import convert

from convert import *
from pdf_handle import *

#from pydrive.auth import GoogleAuth
#from pydrive.drive import GoogleDrive

fpath = "./data_66/"
pdfpath = "./"
fname = "data_66_01_03"
#fname = "data_63_09_26ชุดที่2"
date = ""

def main(_fname):
  global fpath, fname, date

  fname = _fname
  f_tmp1 = open(str(fpath+"database.tmp1"), "w", encoding="utf8")
  f_tmp1.close()

  date = get_date()
  if not total_expand(date):
    exit()
  
  # copy data in database.txt to database.tmp1
  f_database = open(str(fpath+"database.txt"), "r", encoding="utf8")
  f_tmp3 = open(str(fpath+"database.tmp1"), "a", encoding="utf8")
  for x in f_database.readlines():
    f_tmp3.writelines(x)
  f_database.close()
  f_tmp3.close()
  
  seq = get_seq()
  information = read_info()

  for x in information:
    print(x.__dict__)

  information = processing(information)

  print('-----')
  for i in range(len(information)):
    information[i].total = get_old_total(str(fpath+"database.tmp1"), information[i].id)

  print('-----')
  for x in information:
    print(x.__dict__)
    
  document = Document()

  # page setup to A4
  section = document.sections[0]
  section.page_height = Mm(297)
  section.page_width = Mm(210)
  section.left_margin = Inches(1)
  section.right_margin = Inches(1)
  section.top_margin = Inches(1)
  section.bottom_margin = Inches(0.7)

  # Default style
  style = document.styles['Normal']
  font = style.font
  font.name = 'TH SarabunPSK'
  font.size = Pt(36)
  
  styles = document.styles
  style_a36 = styles.add_style('A36', WD_STYLE_TYPE.PARAGRAPH)
  style_a36.font.name = 'TH SarabunPSK'
  style_a36.font.size = Pt(36)
  style_a36b = styles.add_style('A36B', WD_STYLE_TYPE.PARAGRAPH)
  style_a36b.font.name = 'TH SarabunPSK'
  style_a36b.font.size = Pt(36)
  style_a36b.font.bold = True
  style_a36b.font.underline = True

  style_a33 = styles.add_style('A33', WD_STYLE_TYPE.PARAGRAPH)
  style_a33.font.name = 'TH SarabunPSK'
  style_a33.font.size = Pt(33)
  style_a33b = styles.add_style('A33B', WD_STYLE_TYPE.PARAGRAPH)
  style_a33b.font.name = 'TH SarabunPSK'
  style_a33b.font.size = Pt(33)
  style_a33b.font.bold = True
  style_a33b.font.underline = True

  style_a30 = styles.add_style('A30', WD_STYLE_TYPE.PARAGRAPH)
  style_a30.font.name = 'TH SarabunPSK'
  style_a30.font.size = Pt(30)
  style_a30b = styles.add_style('A30B', WD_STYLE_TYPE.PARAGRAPH)
  style_a30b.font.name = 'TH SarabunPSK'
  style_a30b.font.size = Pt(30)
  style_a30b.font.bold = True
  style_a30b.font.underline = True

  style_a27 = styles.add_style('A27', WD_STYLE_TYPE.PARAGRAPH)
  style_a27.font.name = 'TH SarabunPSK'
  style_a27.font.size = Pt(27)
  style_a27b = styles.add_style('A27B', WD_STYLE_TYPE.PARAGRAPH)
  style_a27b.font.name = 'TH SarabunPSK'
  style_a27b.font.size = Pt(27)
  style_a27b.font.bold = True
  style_a27b.font.underline = True

  style_a26 = styles.add_style('A26', WD_STYLE_TYPE.PARAGRAPH)
  style_a26.font.name = 'TH SarabunPSK'
  style_a26.font.size = Pt(25)
  style_a26b = styles.add_style('A26B', WD_STYLE_TYPE.PARAGRAPH)
  style_a26b.font.name = 'TH SarabunPSK'
  style_a26b.font.size = Pt(25)
  style_a26b.font.bold = True
  style_a26b.font.underline = True

  style_a24 = styles.add_style('A24', WD_STYLE_TYPE.PARAGRAPH)
  style_a24.font.name = 'TH SarabunPSK'
  style_a24.font.size = Pt(24)
  style_a24b = styles.add_style('A24B', WD_STYLE_TYPE.PARAGRAPH)
  style_a24b.font.name = 'TH SarabunPSK'
  style_a24b.font.size = Pt(24)
  style_a24b.font.bold = True
  style_a24b.font.underline = True

  style_a22 = styles.add_style('A22', WD_STYLE_TYPE.PARAGRAPH)
  style_a22.font.name = 'TH SarabunPSK'
  style_a22.font.size = Pt(22)
  style_a22b = styles.add_style('A22B', WD_STYLE_TYPE.PARAGRAPH)
  style_a22b.font.name = 'TH SarabunPSK'
  style_a22b.font.size = Pt(22)
  style_a22b.font.bold = True
  style_a22b.font.underline = True

  style_a20 = styles.add_style('A20', WD_STYLE_TYPE.PARAGRAPH)
  style_a20.font.name = 'TH SarabunPSK'
  style_a20.font.size = Pt(20)
  style_a20b = styles.add_style('A20B', WD_STYLE_TYPE.PARAGRAPH)
  style_a20b.font.name = 'TH SarabunPSK'
  style_a20b.font.size = Pt(20)
  style_a20b.font.bold = True
  style_a20b.font.underline = True

  style_a19 = styles.add_style('A19', WD_STYLE_TYPE.PARAGRAPH)
  style_a19.font.name = 'TH SarabunPSK'
  style_a19.font.size = Pt(10)
  style_a19b = styles.add_style('A19B', WD_STYLE_TYPE.PARAGRAPH)
  style_a19b.font.name = 'TH SarabunPSK'
  style_a19b.font.size = Pt(10)
  style_a19b.font.bold = True
  style_a19b.font.underline = True

  f_tmp2 = open(fpath+"database.tmp2", "w", encoding="utf8")
  f_tmp2.writelines("DATE : "+date+'\n')

  for j in range(len(information)):
    information[j].info = insert_space(information[j].info)

  font_sz = [36, 33, 30, 27, 24, 22, 20, 19, 18]
  lines   = [ 6,  7,  8,  9, 10, 11, 12, 14, 15]
  charact = [16, 16, 18, 20, 22, 24]

  j=0
  while j<len(information):
    line=0
    if information[j].id_ex != "":
      line=line+1
    if information[j].count != 0.0:
      line=line+1
      if not (information[j].total[0] == '='): # default value of total is '='
        line=line+1
    if information[j].cash != 0.0:
      line=line+1

    if len(information[j].info)+line > 6:
      while reduce_line(information[j].info):
        print(information[j].info)
    information[j].line = len(information[j].info)+line
    print("ID:%s, line:%d, font:%d, font_indx:%d" % (information[j].id, information[j].line, information[j].font_sz, max(0, information[j].line-6)))
    information[j].font_sz = font_sz[min(8, max(0, information[j].line-6))]
    j=j+1
  
  seq=seq+1
  for j in range(len(information)):
    print("ID:%s, line:%d, font:%d" % (information[j].id, information[j].line, information[j].font_sz))

    style = document.styles[str('A'+str(information[j].font_sz))]
    styleB = document.styles[str('A'+str(information[j].font_sz)+'B')]
    p = document.add_paragraph('No.', style)

    p.add_run(str(seq))
    f_tmp2.write(str(seq)+",\t")
    p = document.add_paragraph('', style)
    _id = information[j].id[1:]
    if _id.isnumeric():
      p.add_run('สมาชิก : ')
      #f_tmp2.write("L")
    p.add_run(str(information[j].id))
    print("%d : %s" % (len(information[j].id), count_space(information[j].id)))
    f_tmp2.write(str(information[j].id)+"\t")
    if information[j].id_ex != "":
      p = document.add_paragraph(str(information[j].id_ex), style)
    p = document.add_paragraph(str('วันที่ : ')+str(date), style)
    p = document.add_paragraph(str('----------'), style)
    
    for k in range(len(information[j].info)):
      p = document.add_paragraph(str(information[j].info[k]), style)

    p = document.add_paragraph(str('----------'), style)
    if information[j].count != 0.0:
      p = document.add_paragraph(str('รวม ')+str(information[j].count)+str(' ชิ้น'), style)
      if not (information[j].total[0] == '='):
        print(information[j].__dict__)
        old_total = float(information[j].total)
        new_total = old_total-information[j].count
        if new_total == int(new_total):
          new_total = int(new_total)
        print("%s, count:%d, old:%f, new:%f" % (information[j].id, information[j].count, old_total, new_total))
        if new_total < 10:
          p = document.add_paragraph(str('เหลือ ')+str(new_total)+str(' ชิ้น'), styleB)
        else:
          p = document.add_paragraph(str('เหลือ ')+str(new_total)+str(' ชิ้น'), style)
        f_tmp2.write("\t,total:"+str(new_total))
    if information[j].cash != 0.0:
      p = document.add_paragraph(str('ยอดชำระ ')+str(information[j].cash)+str(' บาท'), styleB)
      f_tmp2.write("\t,cash:"+str(information[j].cash))
    f_tmp2.writelines("\n")
    line=information[j].line
    while line < 6:
      p = document.add_paragraph('', style)
      line=line+1
    seq=seq+1

  docx = pdfpath+'demo_'+fname+'.docx'
  print("docx="+docx)
  pdf = pdfpath+'demo_'+fname+'.pdf'
  print("pdf="+pdf)
  document.save(docx)
  
  Convert().DOCXtoPDF(docx, pdf)
  #convert(docx, pdf)
  
  #reorganize_pdf(pdf, pdf.replace("demo_data", "reformat"), 6)
  
  #document.save('debug.docx')

#  _database = open(fpath+"database.txt", "r", encoding="utf8")
  f_tmp1 = open(fpath+"database.tmp1", "r", encoding="utf8")
  for t in f_tmp1.readlines():
    f_tmp2.write(t)
#  for t in f_database.readlines():
#    f_tmp2.write(t)
  f_tmp2.close()
  f_tmp1.close()

  os.remove(fpath+"database.old")
  os.rename(fpath+"database.txt", fpath+"database.old")
  os.rename(fpath+"database.tmp2", fpath+"database.txt")

#  f_database.close()

##############################
##############################
##############################

#  gauth = GoogleAuth()
#  gauth.LocalWebserverAuth() # client_secrets.json need to be in the same directory as the script
#  drive = GoogleDrive(gauth)
  
#  folder_id = '1fKs5geKxK3i3aNMvrMTW6ZIfuoganM-N'
#  file1 = drive.CreateFile({"parents": [{"kind": "drive#fileLink", "id": folder_id}]})
#  file1.SetContentFile(docx)
#  file1.Upload() # Upload the file.
  
  # View all folders and file in your Google Drive
#  fileList = drive.ListFile({'q': "'{0}' in parents and trashed=false".format(folder_id)}).GetList()
#  for file in fileList:
#    if file['title'] == docx:
#      print('Title: %s, ID: %s' % (file['title'], file['id']))

def get_date():
  global fname
  f_data = open(str(fpath+fname+".txt"), "r", encoding="utf8")
  _date = f_data.readline()
  f_data.close()
  if _date[len(_date)-1] == '\n':
    _date = _date[:len(_date)-1]
  if _date.find("/64") == -1:
    print("Invalid date ...")
  return _date

def total_expand(_date):
  global fname
  f_data = open(str(fpath+fname+".txt"), "r", encoding="utf8")
  f_tmp1 = open(str(fpath+"database.tmp1"), "a", encoding="utf8")

  i=0  
  for x in f_data.readlines():
    txt = x.split()
    print(txt)
    if txt[0] == "++":
      _id = txt[1]
      _value = float(txt[2])
      _old_total = get_old_total(str(fpath+"database.txt"), "L"+_id)
      if _old_total != '-':
        _new_total = float(_old_total) + _value
        _new_total = float_to_string(_new_total)
        print("line[%d]: %s, %d ::: old, new : %f, %s" % (i, _id, _value, float(_old_total), _new_total))
        new_txt = [str("SPECIAL : "+_date+", L"+_id+"  ต่อสมาชิก "+str(_value)+" ชิ้น, total:"+_new_total+'\n')]
        #print(new_txt)
        f_tmp1.writelines(new_txt)
      else:
        print(txt)
        return False
    else:
      print("txt[%d] == ++ is FALSE : %s" % (i, txt))
    i=i+1
  
  f_data.close()
  f_tmp1.close()
  return True

def get_seq():
  f = open(str(fpath+"database.txt"), "r", encoding="utf8")
  _max = 0
  for i in range(30):
    txt = f.readline()
    if len(txt) == 0:
      continue
    txt = txt.split()
    if len(txt) == 0:
      continue
    txt = txt[0]
    txt = txt[:len(txt)-1]
    if txt.isnumeric():
      if _max < int(txt):
        if int(txt) > 1000:
          break
        _max = int(txt)
  f.close()
  return _max

def read_info():
  global fname
  f = open(str(fpath+fname+".txt"), "r", encoding="utf8")
  lines = f.readlines()
  f.close()
  data = []
  for line in lines:
    if line.find("++") == -1:
      if line[len(line)-1] == '\n':
        line = line[:len(line)-1]
      data.append(line)
      print(line)


  #date = data[0]
  #seq = int(seq)
  #print(date)
  #print(seq)

  information = []
  for i in range(1, len(data)):
    tmp = data[i].split()
    print(tmp)
    v = Person()
    v.id = tmp[0]
    if v.id.isnumeric():
      v.id = str("L" + v.id)
    start = 1
    if tmp[1][0] == '(':
      v.id_ex = tmp[1]
      start = 2
    v.info = tmp[start:len(tmp)]
    #v.info = tmp[1:len(tmp)-1]
    #v.total = tmp[len(tmp)-1:]
    information.append(v)

  return information

def get_old_total(_fname, _id):
  #print("get_old_total, name:%s, _id:%s, type:%s" % (name, _id, type(_id)))
  f = open(_fname, "r", encoding="utf8")
  b_found=False
  lines = f.readlines()
  f.close()
  for line in lines:
    #print(line)
    indx = line.find(_id)
    if indx != -1:
      lstr = line.split(",")
      s_total=""
      for y in lstr:
        indx_total=y.find("total:")
        if indx_total != -1:
          b_found=True
          s_total=y[indx_total+6:len(y)-1]
          #print("%s" % s_total)
          return s_total
  if not b_found:
    print("not_found ID:%s" % _id)
    return "-"

def processing(information):
  for v in information:
    _b_err = False
    _sum_cash = 0.0
    print(v.info)
    for k in range(len(v.info)):
      i=0
      if v.info[k].find('*') != -1:
        while i < len(v.info[k]):
          if v.info[k][i].isnumeric():
            break
          i=i+1
        if i != 0:
          new_info = split(["*","="], v.info[k][i:])
          new_info = [float(i) for i in new_info]
          mul = new_info[0]*new_info[1]
          print("new=%s : old=%s : %s" %(new_info, v.info[k], normal_round(new_info[0]*new_info[1])))
          if new_info[1] > 4:
            mul = normal_round(mul)
            v.cash = v.cash+mul
            #print("%s | %s" % (mul, _sum_cash))
          else:
            v.count = v.count + mul
          if (len(new_info) == 3) and (mul != new_info[2]):
            print("new=%s : old=%s : %s" %(new_info, _info, normal_round(new_info[0]*new_info[1])))
            _b_err = True
            break
          elif len(new_info) == 2:
            #print('In')
            #print(v.info[k])
            v.info[k] = v.info[k]+str('=')+str(mul)
            #print(v.info[k])
            #print('--In--')
      elif not (v.info[k][0] == 'ม'):
        #print("debug : %s" % v.info[k])
        i=0
        while i < len(v.info[k]):
          if v.info[k][i].isnumeric():
            break
          i=i+1
        if i != 0:
          print("%s : v.info[%d][%d:] : %s" % (v.info, k, i, v.info[k][i:]))
          new_info = int(v.info[k][i:])
          v.count = v.count + new_info
    if v.count == int(v.count):
      v.count = int(v.count)
    if v.cash == int(v.cash):
      v.cash = int(v.cash)
      #print("v.count, int(v.count) : %f, %d" % (v.count, int(v.count)))
    #print(v.info)
            
    if _b_err:
      print("ERROR : %s" % v.__dict__)
  return information

def insert_space(info):
  for k in range(len(info)):
    i=0
    while i < len(info[k]):
      if info[k][i].isnumeric():
        break
      i=i+1
    if i != 0:
      info[k] = info[k][:i] + str(' ') + info[k][i:]
  return info

def reduce_line(info):  
  '''
  for i in range(len(info)-1):
    for j in range(i+1, len(info)):
      x=info[i]
      y=info[j]
      if len(x+"    "+y) < 20:
        info[i] = info[i]+"    "+info[j]
        del info[j]
        return True
  return False
  '''
  i=0
  while i < len(info)-1:
    indx_del = []
    s = info[i]
    for j in range(i+1, len(info)):
      y = info[j]
      if len(s + "  " + y) < 20:
        s = s + "  " + y
        indx_del.append(j)
      else:
        break
    if len(indx_del) != 0:
      info[i] = s
      print('indx_del')
      for k in reversed(range(len(indx_del))):
        print(info[indx_del[k]])
        del info[indx_del[k]]
      print('indx_del - end')
      indx_del = []
    i+=1

def count_space(txt):
  _sum = len(txt)
  thai_vowels = ['ู', 'ุ', 'ึ', 'ั', '๊', 'ี', '็', '้', '๋', '่', 'ิ', '์', 'ื']
  #print("thai_vowels : %d : type: %s" % ord(thai_vowels))
  for x in txt:
    is_found = False
    for y in thai_vowels:
      if ord(x) == ord(y):
        is_found = True
        break
    if is_found:
      _sum = _sum-1
      #print("count_space : %s" % txt)
  return _sum  

class Person:
  def __init__(self):
    self.id = ""
    self.id_ex = ""
    self.info = []
    self.total = '='
    self.count = 0.0
    self.cash = 0.0
    self.old_total = 0.0
    self.font_sz = 36
    self.line = 6

def split(delimiters, string, maxsplit=0):
    import re
    regexPattern = '|'.join(map(re.escape, delimiters))
    return re.split(regexPattern, string, maxsplit)

def normal_round(n):
    if n - math.floor(n) < 0.5:
        return math.floor(n)
    return math.ceil(n)

def float_to_string(_num):
  if _num == int(_num):
    _num = int(_num)
  return str(_num)

if __name__ == '__main__':  
  print('Number of arguments: ' + str(len(sys.argv)))
  print('Argument List: ' + str(sys.argv))
  
  if len(sys.argv) > 1:
    if sys.argv[1].find('data_66/data_66_') != -1:
      _fname = sys.argv[1][8:-4]
      print("fname: " + _fname)
      main(_fname)
    else:
      print("if sys.argv[1].find('data_66/data_66') != -1 is FALSE")
  else:
    print("if len(sys.argv) > 1 is FALSE")


